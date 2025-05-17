import os
import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.cluster import KMeans
from sklearn.linear_model import LinearRegression
from sklearn.model_selection import train_test_split
from sklearn.preprocessing import LabelEncoder
from docx import Document
from docx.shared import Inches

# # Try docx2pdf for optional PDF export
# try:
#     from docx2pdf import convert
#     PDF_CONVERT = True
# except ImportError:
#     PDF_CONVERT = False

st.set_page_config(page_title="AI Data Analyst", layout="wide")

st.title("🤖 AI Data Analyst")
st.markdown("Upload your dataset and generate a detailed Word (and optional PDF) report with charts and insights.")

# Upload data
uploaded_file = st.file_uploader("Upload your data file (CSV, Excel, JSON)", type=["csv", "xlsx", "xls", "json"])

def load_any_file(uploaded_file):
    try:
        if uploaded_file.name.endswith('.csv'):
            return pd.read_csv(uploaded_file)
        elif uploaded_file.name.endswith('.json'):
            return pd.read_json(uploaded_file)
        elif uploaded_file.name.endswith(('.xls', '.xlsx')):
            return pd.read_excel(uploaded_file)
        else:
            st.error("Unsupported file format.")
            return None
    except Exception as e:
        st.error(f"Error loading file: {e}")
        return None

def clean_data(df):
    df.columns = df.columns.str.strip().str.lower().str.replace(' ', '_')
    df = df.drop_duplicates()
    df = df.fillna(method='ffill').fillna(method='bfill')
    return df

def summary_statistics(df):
    st.subheader("📊 Summary Statistics")
    st.write(df.describe())
    corr = df.corr(numeric_only=True)
    st.subheader("🔗 Correlation Heatmap")
    fig, ax = plt.subplots()
    sns.heatmap(corr, annot=True, cmap="coolwarm", ax=ax)
    plt.title("Correlation Heatmap")
    fig.tight_layout()
    st.pyplot(fig)
    fig.savefig("output/correlation_heatmap.png")

def run_kmeans(df, n_clusters=3):
    df_numeric = df.select_dtypes(include='number').dropna()
    model = KMeans(n_clusters=n_clusters, random_state=0)
    df['cluster'] = model.fit_predict(df_numeric)
    return df

def run_linear_regression(df, target_col):
    df = df.dropna()
    if target_col not in df.columns:
        return None
    df_encoded = df.copy()
    for col in df.columns:
        if df[col].dtype == 'object':
            df_encoded[col] = LabelEncoder().fit_transform(df[col].astype(str))
    X = df_encoded.drop(columns=[target_col])
    y = df_encoded[target_col]
    X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=0)
    model = LinearRegression()
    model.fit(X_train, y_train)
    score = model.score(X_test, y_test)
    return {
        "target": target_col,
        "r2": score,
        "coefficients": dict(zip(X.columns, model.coef_))
    }

def generate_detailed_summary(df, regression_results=None):
    summary = []
    summary.append(f"Data shape: {df.shape[0]} rows, {df.shape[1]} columns.\n")
    summary.append("### Columns and Types:\n" + str(df.dtypes))
    missing = df.isnull().sum()
    missing = missing[missing > 0]
    if not missing.empty:
        summary.append("\n### Missing Values:\n" + str(missing))
    else:
        summary.append("\nNo missing values.\n")
    desc = df.describe().round(2).to_string()
    summary.append("\n### Descriptive Statistics:\n" + desc)
    if regression_results:
        summary.append("\n### Regression Results:")
        summary.append(f"Target: {regression_results['target']}")
        summary.append(f"R² Score: {regression_results['r2']:.3f}")
        summary.append("Coefficients:")
        for k, v in regression_results["coefficients"].items():
            summary.append(f"{k}: {v:.3f}")
    return "\n".join(summary)

def generate_word_report(summary_text, chart_paths, output_path="output/report.docx"):
    doc = Document()
    doc.add_heading('AI Employee Analysis Report', 0)
    for para in summary_text.split('\n'):
        if para.strip().startswith("###"):
            doc.add_heading(para.strip('#').strip(), level=1)
        else:
            doc.add_paragraph(para)
    doc.add_heading('Charts', level=1)
    for chart in chart_paths:
        if os.path.exists(chart):
            doc.add_picture(chart, width=Inches(5.5))
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    # if PDF_CONVERT:
    #     convert(output_path, output_path.replace(".docx", ".pdf"))

# Main flow
if uploaded_file:
    df = load_any_file(uploaded_file)
    if df is not None:
        df = clean_data(df)
        st.success("✅ Data loaded and cleaned!")
        if st.checkbox("Show raw data"):
            st.dataframe(df.head())

        run_summary = st.button("Generate Summary Statistics")
        if run_summary:
            summary_statistics(df)

        if st.button("Run KMeans Clustering"):
            df = run_kmeans(df)
            st.success("✅ KMeans clustering done. Check 'cluster' column in the DataFrame.")

        if st.button("Run Linear Regression"):
            target = st.text_input("Enter the target column for regression:", key="target_col")
            if target:
                result = run_linear_regression(df, target)
                if result:
                    st.write(f"R² Score: {result['r2']:.3f}")
                    st.write(result["coefficients"])
                else:
                    st.error("Invalid target column.")

        if st.button("Generate Report"):
            regression_result = run_linear_regression(df, target) if 'target' in locals() else None
            summary_text = generate_detailed_summary(df, regression_result)
            generate_word_report(summary_text, ["output/correlation_heatmap.png"])
            st.success("📄 Report saved to 'output/report.docx'")
            # if PDF_CONVERT:
            #     st.info("📄 PDF version also available in 'output/report.pdf'")
